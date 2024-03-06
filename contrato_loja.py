from datetime import datetime

from docx import Document
from openpyxl import load_workbook

# Passar as informações que estão na planilha fornecedores para um arquivo word

planilha_clientes = load_workbook('./clientes.xlsx')
pagina_clientes = planilha_clientes['Sheet1']


for linha in pagina_clientes.iter_rows(min_row=2,values_only=True):
    
    nome_cliente, cpf, rg, endereco, bairro, cidade, estado, cep, whatsapp, telefone, email, produto, preco, cupom = linha
    # print(linha)
    arquivo_word = Document()
    arquivo_word.add_heading('Contrato de Aquisição de Produtos')
      
    
    texto_contrato = f""" 
Partes Contratantes:

De um lado, como VENDEDOR, a empresa:

Loja de Variedades Concipa - LTDA
CNPJ: 00.000.000.000-00
Endereço: Rua dos vendedores, 900, loja 5, Centro - Conceição de Ipanema - MG
CEP: 36947-000
E-mail: elciodev@gmail.com
Telefone: (11) 9 1432-7845
WhatsApp:(11) 9 1432-7845

E de outro lado, como COMPRADOR, Sr(a). {nome_cliente}, portador(a) do CPF: {cpf} e RG: {rg}, residente e domiciliado(a) em {endereco}, {bairro} - {cidade} - {estado}, CEP: {cep}, doravante denominado simplesmente COMPRADOR.

Objeto:

O presente contrato tem como objeto a aquisição do seguinte produto:

Produto: {produto}

Valor: R$ {preco}

Número da Nota Fiscal/Cupom: {cupom}

Prazo de Pagamento:

O pagamento deverá ser efetuado em até 30 (trinta) dias a partir da data da emissão da presente nota fiscal, através de boleto bancário.

Multa por Atraso:

Em caso de atraso no pagamento, o COMPRADOR ficará sujeito a uma multa de 5% (cinco por cento) sobre o valor total do produto.



Juros por Dia de Atraso:

Além da multa estabelecida, o COMPRADOR também pagará juros de 1% (um por cento) ao dia sobre o valor total do produto, referente aos dias de atraso no pagamento.

Disposições Finais:

    A entrega do produto será realizada mediante a assinatura e o reconhecimento deste documento comprometendo-se o comprador com a quitação integral do valor acordado.

    Em caso de atraso no pagamento por parte do COMPRADOR, o VENDEDOR poderá rescindir o contrato, ficando autorizado a buscar as medidas legais cabíveis para a cobrança do débito entre elas inclusão nos serviços de proteção ao crédito, Serasa e Protesto.

    Qualquer alteração neste contrato só terá validade se efetuada por escrito e assinada por ambas as partes.

E, por estarem assim justas e contratadas, firmam o presente em duas vias, juntamente com 2 (duas) testemunhas.

Local e Data:

{cidade}, {datetime.now().strftime('%d/%m/%Y')}

Loja de Variedades Concipa - LTDA

CNPJ: 00.000.000.000-00

Comprador: {nome_cliente}

CPF: {cpf}
RG: {rg}

Testemunhas:

Nome:___________________________________________________

CPF:_______________


Nome:___________________________________________________

CPF:_______________
    """
      
    arquivo_word.add_paragraph(texto_contrato)
    # 
    # Salvando arquivo em uma pasta espeífica
    arquivo_word.save(f'./contratos/contrato_{nome_cliente}.docx')